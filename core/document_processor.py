"""
Document Processing Engine for RAG System
Handles parsing, chunking, and preprocessing of various document types
"""

import os
import re
import hashlib
import asyncio
import logging
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
from dataclasses import dataclass
import mimetypes

# Document parsers
import PyMuPDF as fitz  # PDF processing
from docx import Document as DocxDocument  # Word documents
import pandas as pd
from bs4 import BeautifulSoup
import markdown

# Code analysis
import ast
import tokenize
from io import StringIO
import tree_sitter_python as tspython
import tree_sitter_c as tsc
import tree_sitter_cpp as tscpp
import tree_sitter
from pygments import highlight
from pygments.lexers import get_lexer_for_filename, guess_lexer
from pygments.formatters import NullFormatter

@dataclass
class ProcessedDocument:
    """Container for processed document data"""
    file_path: str
    file_name: str
    file_type: str
    file_size: int
    content_hash: str
    security_classification: str
    domain: str
    language: str
    chunks: List[Dict[str, Any]]
    metadata: Dict[str, Any]

@dataclass
class DocumentChunk:
    """Individual document chunk"""
    content: str
    content_type: str  # 'text', 'code', 'table', 'image_caption'
    chunk_index: int
    start_char: int
    end_char: int
    metadata: Dict[str, Any]

class DocumentProcessor:
    """Main document processing engine"""
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.logger = logging.getLogger(__name__)
        
        # Initialize tree-sitter parsers
        self.parsers = self._init_code_parsers()
        
        # Supported file types and their processors
        self.processors = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_doc,
            '.odt': self._process_odt,
            '.txt': self._process_text,
            '.md': self._process_markdown,
            '.rst': self._process_rst,
            '.c': self._process_c_code,
            '.h': self._process_c_code,
            '.cpp': self._process_cpp_code,
            '.cxx': self._process_cpp_code,
            '.cc': self._process_cpp_code,
            '.hpp': self._process_cpp_code,
            '.hxx': self._process_cpp_code,
            '.py': self._process_python_code,
            '.m': self._process_matlab_code,
            '.html': self._process_html,
            '.xml': self._process_xml,
            '.json': self._process_json,
            '.csv': self._process_csv,
            '.xlsx': self._process_excel,
            '.sql': self._process_sql,
        }
        
        # Chunk size configurations
        self.chunk_configs = {
            'text': {
                'size': config.get('embedding', {}).get('chunk_size', 1000),
                'overlap': config.get('embedding', {}).get('chunk_overlap', 200)
            },
            'code': {
                'size': 800,  # Smaller chunks for code
                'overlap': 100
            },
            'table': {
                'size': 1500,  # Larger chunks for tables
                'overlap': 0
            }
        }
    
    def _init_code_parsers(self) -> Dict[str, tree_sitter.Parser]:
        """Initialize tree-sitter parsers for code analysis"""
        parsers = {}
        
        try:
            # Python parser
            python_parser = tree_sitter.Parser()
            python_parser.set_language(tree_sitter.Language(tspython.language(), "python"))
            parsers['python'] = python_parser
            
            # C parser
            c_parser = tree_sitter.Parser()
            c_parser.set_language(tree_sitter.Language(tsc.language(), "c"))
            parsers['c'] = c_parser
            
            # C++ parser
            cpp_parser = tree_sitter.Parser()
            cpp_parser.set_language(tree_sitter.Language(tscpp.language(), "cpp"))
            parsers['cpp'] = cpp_parser
            
            self.logger.info(f"Initialized {len(parsers)} code parsers")
            
        except Exception as e:
            self.logger.warning(f"Error initializing code parsers: {e}")
        
        return parsers
    
    async def process_file(self, 
                          file_path: str, 
                          security_classification: str = "internal",
                          domain: str = "general") -> Optional[ProcessedDocument]:
        """Process a single file"""
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                self.logger.error(f"File not found: {file_path}")
                return None
            
            # Get file info
            file_stats = file_path.stat()
            file_ext = file_path.suffix.lower()
            file_type = self._determine_file_type(file_path)
            language = self._determine_language(file_path)
            
            # Calculate content hash
            content_hash = self._calculate_file_hash(file_path)
            
            # Check if we have a processor for this file type
            if file_ext not in self.processors:
                self.logger.warning(f"No processor for file type: {file_ext}")
                return None
            
            # Process the file
            chunks = await self.processors[file_ext](file_path)
            
            if not chunks:
                self.logger.warning(f"No content extracted from: {file_path}")
                return None
            
            # Create processed document
            processed_doc = ProcessedDocument(
                file_path=str(file_path),
                file_name=file_path.name,
                file_type=file_type,
                file_size=file_stats.st_size,
                content_hash=content_hash,
                security_classification=security_classification,
                domain=domain,
                language=language,
                chunks=chunks,
                metadata={
                    'created_time': file_stats.st_ctime,
                    'modified_time': file_stats.st_mtime,
                    'extension': file_ext,
                    'encoding': 'utf-8',
                    'total_chunks': len(chunks),
                    'total_chars': sum(len(chunk['content']) for chunk in chunks)
                }
            )
            
            self.logger.info(f"Processed {file_path}: {len(chunks)} chunks, "
                           f"{processed_doc.metadata['total_chars']} characters")
            
            return processed_doc
            
        except Exception as e:
            self.logger.error(f"Error processing file {file_path}: {e}")
            return None
    
    def _determine_file_type(self, file_path: Path) -> str:
        """Determine the type of file"""
        ext = file_path.suffix.lower()
        
        if ext in ['.c', '.h', '.cpp', '.cxx', '.cc', '.hpp', '.hxx']:
            return 'source_code'
        elif ext in ['.py', '.m']:
            return 'source_code'
        elif ext in ['.pdf', '.docx', '.doc', '.odt']:
            return 'document'
        elif ext in ['.txt', '.md', '.rst']:
            return 'text'
        elif ext in ['.html', '.xml']:
            return 'markup'
        elif ext in ['.csv', '.xlsx']:
            return 'data'
        elif ext in ['.json', '.yaml', '.yml']:
            return 'config'
        else:
            return 'unknown'
    
    def _determine_language(self, file_path: Path) -> str:
        """Determine the programming language"""
        ext = file_path.suffix.lower()
        
        language_map = {
            '.c': 'c',
            '.h': 'c',
            '.cpp': 'cpp',
            '.cxx': 'cpp',
            '.cc': 'cpp',
            '.hpp': 'cpp',
            '.hxx': 'cpp',
            '.py': 'python',
            '.m': 'matlab',
            '.html': 'html',
            '.xml': 'xml',
            '.js': 'javascript',
            '.sql': 'sql'
        }
        
        return language_map.get(ext, 'text')
    
    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate MD5 hash of file content"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    
    # Document processors for different file types
    
    async def _process_pdf(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process PDF documents"""
        chunks = []
        
        try:
            doc = fitz.open(file_path)
            full_text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                
                if text.strip():
                    full_text += f"\n--- Page {page_num + 1} ---\n{text}"
            
            doc.close()
            
            if full_text.strip():
                # Chunk the text
                text_chunks = self._chunk_text(full_text, 'text')
                
                for i, chunk_text in enumerate(text_chunks):
                    chunks.append({
                        'content': chunk_text,
                        'content_type': 'text',
                        'chunk_index': i,
                        'metadata': {
                            'source_type': 'pdf',
                            'total_pages': len(doc) if 'doc' in locals() else 0
                        }
                    })
            
        except Exception as e:
            self.logger.error(f"Error processing PDF {file_path}: {e}")
        
        return chunks
    
    async def _process_docx(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process Word DOCX documents"""
        chunks = []
        
        try:
            doc = DocxDocument(file_path)
            full_text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                table_text = "\n--- Table ---\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    table_text += row_text + "\n"
                full_text += table_text + "\n"
            
            if full_text.strip():
                # Chunk the text
                text_chunks = self._chunk_text(full_text, 'text')
                
                for i, chunk_text in enumerate(text_chunks):
                    chunks.append({
                        'content': chunk_text,
                        'content_type': 'text',
                        'chunk_index': i,
                        'metadata': {
                            'source_type': 'docx',
                            'has_tables': len(doc.tables) > 0
                        }
                    })
            
        except Exception as e:
            self.logger.error(f"Error processing DOCX {file_path}: {e}")
        
        return chunks
    
    async def _process_doc(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process legacy Word DOC documents using LibreOffice"""
        chunks = []
        
        try:
            # Convert DOC to text using LibreOffice
            import subprocess
            import tempfile
            
            with tempfile.TemporaryDirectory() as temp_dir:
                # Convert to text
                result = subprocess.run([
                    'libreoffice', '--headless', '--convert-to', 'txt',
                    '--outdir', temp_dir, str(file_path)
                ], capture_output=True, text=True)
                
                if result.returncode == 0:
                    txt_file = Path(temp_dir) / f"{file_path.stem}.txt"
                    if txt_file.exists():
                        with open(txt_file, 'r', encoding='utf-8') as f:
                            content = f.read()
                        
                        if content.strip():
                            text_chunks = self._chunk_text(content, 'text')
                            
                            for i, chunk_text in enumerate(text_chunks):
                                chunks.append({
                                    'content': chunk_text,
                                    'content_type': 'text',
                                    'chunk_index': i,
                                    'metadata': {'source_type': 'doc'}
                                })
                else:
                    self.logger.warning(f"LibreOffice conversion failed for {file_path}")
            
        except Exception as e:
            self.logger.error(f"Error processing DOC {file_path}: {e}")
        
        return chunks
    
    async def _process_odt(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process OpenDocument Text files"""
        chunks = []
        
        try:
            # Use LibreOffice to convert ODT to text
            import subprocess
            import tempfile
            
            with tempfile.TemporaryDirectory() as temp_dir:
                result = subprocess.run([
                    'libreoffice', '--headless', '--convert-to', 'txt',
                    '--outdir', temp_dir, str(file_path)
                ], capture_output=True, text=True)
                
                if result.returncode == 0:
                    txt_file = Path(temp_dir) / f"{file_path.stem}.txt"
                    if txt_file.exists():
                        with open(txt_file, 'r', encoding='utf-8') as f:
                            content = f.read()
                        
                        if content.strip():
                            text_chunks = self._chunk_text(content, 'text')
                            
                            for i, chunk_text in enumerate(text_chunks):
                                chunks.append({
                                    'content': chunk_text,
                                    'content_type': 'text',
                                    'chunk_index': i,
                                    'metadata': {'source_type': 'odt'}
                                })
            
        except Exception as e:
            self.logger.error(f"Error processing ODT {file_path}: {e}")
        
        return chunks
    
    async def _process_text(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process plain text files"""
        chunks = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if content.strip():
                text_chunks = self._chunk_text(content, 'text')
                
                for i, chunk_text in enumerate(text_chunks):
                    chunks.append({
                        'content': chunk_text,
                        'content_type': 'text',
                        'chunk_index': i,
                        'metadata': {'source_type': 'text'}
                    })
        
        except Exception as e:
            self.logger.error(f"Error processing text file {file_path}: {e}")
        
        return chunks
    
    async def _process_markdown(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process Markdown files"""
        chunks = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Convert markdown to HTML then to text for better structure
            html = markdown.markdown(content)
            soup = BeautifulSoup(html, 'html.parser')
            text_content = soup.get_text()
            
            if text_content.strip():
                text_chunks = self._chunk_text(text_content, 'text')
                
                for i, chunk_text in enumerate(text_chunks):
                    chunks.append({
                        'content': chunk_text,
                        'content_type': 'text',
                        'chunk_index': i,
                        'metadata': {
                            'source_type': 'markdown',
                            'original_markdown': True
                        }
                    })
        
        except Exception as e:
            self.logger.error(f"Error processing Markdown {file_path}: {e}")
        
        return chunks
    
    async def _process_rst(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process reStructuredText files"""
        chunks = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # For now, treat as plain text
            # TODO: Add proper RST parsing
            if content.strip():
                text_chunks = self._chunk_text(content, 'text')
                
                for i, chunk_text in enumerate(text_chunks):
                    chunks.append({
                        'content': chunk_text,
                        'content_type': 'text',
                        'chunk_index': i,
                        'metadata': {'source_type': 'rst'}
                    })
        
        except Exception as e:
            self.logger.error(f"Error processing RST {file_path}: {e}")
        
        return chunks
    
    async def _process_c_code(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process C source code files"""
        chunks = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Use tree-sitter for semantic chunking
            if 'c' in self.parsers:
                chunks.extend(self._chunk_code_semantically(content, 'c'))
            else:
                # Fallback to simple chunking
                code_chunks = self._chunk_text(content, 'code')
                for i, chunk_text in enumerate(code_chunks):
                    chunks.append({
                        'content': chunk_text,
                        'content_type': 'code',
                        'chunk_index': i,
                        'metadata': {
                            'language': 'c',
                            'file_type': file_path.suffix
                        }
                    })
        
        except Exception as e:
            self.logger.error(f"Error processing C code {file_path}: {e}")
        
        return chunks
    
    async def _process_cpp_code(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process C++ source code files"""
        chunks = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Use tree-sitter for semantic chunking
            if 'cpp' in self.parsers:
                chunks.extend(self._chunk_code_semantically(content, 'cpp'))
            else:
                # Fallback to simple chunking
                code_chunks = self._chunk_text(content, 'code')
                for i, chunk_text in enumerate(code_chunks):
                    chunks.append({
                        'content': chunk_text,
                        'content_type': 'code',
                        'chunk_index': i,
                        'metadata': {
                            'language': 'cpp',
                            'file_type': file_path.suffix
                        }
                    })
        
        except Exception as e:
            self.logger.error(f"Error processing C++ code {file_path}: {e}")
        
        return chunks
    
    async def _process_python_code(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process Python source code files"""
        chunks = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Use tree-sitter for semantic chunking
            if 'python' in self.parsers:
                chunks.extend(self._chunk_code_semantically(content, 'python'))
            else:
                # Fallback to AST-based chunking
                chunks.extend(self._chunk_python_with_ast(content))
        
        except Exception as e:
            self.logger.error(f"Error processing Python code {file_path}: {e}")
        
        return chunks
    
    async def _process_matlab_code(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process MATLAB code files"""
        chunks = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Simple function-based chunking for MATLAB
            function_chunks = self._chunk_matlab_functions(content)
            
            for i, (func_content, func_name) in enumerate(function_chunks):
                chunks.append({
                    'content': func_content,
                    'content_type': 'code',
                    'chunk_index': i,
                    'metadata': {
                        'language': 'matlab',
                        'function_name': func_name,
                        'file_type': '.m'
                    }
                })
        
        except Exception as e:
            self.logger.error(f"Error processing MATLAB code {file_path}: {e}")
        
        return chunks
    
    async def _process_html(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process HTML files"""
        chunks = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            
            # Extract text content
            text_content = soup.get_text()
            if text_content.strip():
                text_chunks = self._chunk_text(text_content, 'text')
                
                for i, chunk_text in enumerate(text_chunks):
                    chunks.append({
                        'content': chunk_text,
                        'content_type': 'text',
                        'chunk_index': i,
                        'metadata': {
                            'source_type': 'html',
                            'has_code': bool(soup.find_all(['code', 'pre']))
                        }
                    })
            
            # Extract code blocks separately
            code_blocks = soup.find_all(['code', 'pre'])
            for i, block in enumerate(code_blocks):
                if block.text.strip():
                    chunks.append({
                        'content': block.text.strip(),
                        'content_type': 'code',
                        'chunk_index': len(chunks),
                        'metadata': {
                            'source_type': 'html_code_block',
                            'tag': block.name
                        }
                    })
        
        except Exception as e:
            self.logger.error(f"Error processing HTML {file_path}: {e}")
        
        return chunks
    
    async def _process_xml(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process XML files"""
        chunks = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            soup = BeautifulSoup(content, 'xml')
            text_content = soup.get_text()
            
            if text_content.strip():
                text_chunks = self._chunk_text(text_content, 'text')
                
                for i, chunk_text in enumerate(text_chunks):
                    chunks.append({
                        'content': chunk_text,
                        'content_type': 'text',
                        'chunk_index': i,
                        'metadata': {'source_type': 'xml'}
                    })
        
        except Exception as e:
            self.logger.error(f"Error processing XML {file_path}: {e}")
        
        return chunks
    
    async def _process_json(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process JSON files"""
        chunks = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert JSON to readable text
            json_text = json.dumps(data, indent=2)
            
            text_chunks = self._chunk_text(json_text, 'text')
            
            for i, chunk_text in enumerate(text_chunks):
                chunks.append({
                    'content': chunk_text,
                    'content_type': 'data',
                    'chunk_index': i,
                    'metadata': {
                        'source_type': 'json',
                        'data_structure': type(data).__name__
                    }
                })
        
        except Exception as e:
            self.logger.error(f"Error processing JSON {file_path}: {e}")
        
        return chunks
    
    async def _process_csv(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process CSV files"""
        chunks = []
        
        try:
            df = pd.read_csv(file_path)
            
            # Create text representation of the CSV
            csv_text = f"CSV File: {file_path.name}\n"
            csv_text += f"Columns: {', '.join(df.columns.tolist())}\n"
            csv_text += f"Rows: {len(df)}\n\n"
            
            # Add sample data
            if len(df) > 0:
                csv_text += "Sample data:\n"
                csv_text += df.head(10).to_string()
            
            # Add column descriptions if available
            csv_text += "\n\nColumn Information:\n"
            for col in df.columns:
                col_info = f"- {col}: {df[col].dtype}"
                if df[col].dtype in ['object', 'string']:
                    unique_vals = df[col].dropna().unique()
                    if len(unique_vals) <= 10:
                        col_info += f" (values: {', '.join(map(str, unique_vals))})"
                csv_text += col_info + "\n"
            
            text_chunks = self._chunk_text(csv_text, 'table')
            
            for i, chunk_text in enumerate(text_chunks):
                chunks.append({
                    'content': chunk_text,
                    'content_type': 'table',
                    'chunk_index': i,
                    'metadata': {
                        'source_type': 'csv',
                        'columns': df.columns.tolist(),
                        'row_count': len(df),
                        'column_types': df.dtypes.to_dict()
                    }
                })
        
        except Exception as e:
            self.logger.error(f"Error processing CSV {file_path}: {e}")
        
        return chunks
    
    async def _process_excel(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process Excel files"""
        chunks = []
        
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Create text representation
                excel_text = f"Excel File: {file_path.name}\n"
                excel_text += f"Sheet: {sheet_name}\n"
                excel_text += f"Columns: {', '.join(df.columns.tolist())}\n"
                excel_text += f"Rows: {len(df)}\n\n"
                
                if len(df) > 0:
                    excel_text += "Sample data:\n"
                    excel_text += df.head(10).to_string()
                
                text_chunks = self._chunk_text(excel_text, 'table')
                
                for i, chunk_text in enumerate(text_chunks):
                    chunks.append({
                        'content': chunk_text,
                        'content_type': 'table',
                        'chunk_index': len(chunks) + i,
                        'metadata': {
                            'source_type': 'excel',
                            'sheet_name': sheet_name,
                            'columns': df.columns.tolist(),
                            'row_count': len(df)
                        }
                    })
        
        except Exception as e:
            self.logger.error(f"Error processing Excel {file_path}: {e}")
        
        return chunks
    
    async def _process_sql(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process SQL files"""
        chunks = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Split SQL into statements
            statements = self._split_sql_statements(content)
            
            for i, statement in enumerate(statements):
                if statement.strip():
                    chunks.append({
                        'content': statement.strip(),
                        'content_type': 'code',
                        'chunk_index': i,
                        'metadata': {
                            'language': 'sql',
                            'statement_type': self._identify_sql_statement_type(statement)
                        }
                    })
        
        except Exception as e:
            self.logger.error(f"Error processing SQL {file_path}: {e}")
        
        return chunks
    
    # Chunking methods
    
    def _chunk_text(self, text: str, content_type: str) -> List[str]:
        """Chunk text based on content type"""
        config = self.chunk_configs.get(content_type, self.chunk_configs['text'])
        chunk_size = config['size']
        overlap = config['overlap']
        
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence ending within the last 200 characters
                sentence_end = max(
                    text.rfind('.', start, end),
                    text.rfind('!', start, end),
                    text.rfind('?', start, end)
                )
                
                if sentence_end > start + chunk_size * 0.5:
                    end = sentence_end + 1
                else:
                    # Look for paragraph break
                    para_break = text.rfind('\n\n', start, end)
                    if para_break > start + chunk_size * 0.3:
                        end = para_break + 2
                    else:
                        # Look for line break
                        line_break = text.rfind('\n', start, end)
                        if line_break > start + chunk_size * 0.5:
                            end = line_break + 1
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
            
            # Prevent infinite loop
            if start <= 0:
                break
        
        return chunks
    
    def _chunk_code_semantically(self, code: str, language: str) -> List[Dict[str, Any]]:
        """Chunk code using tree-sitter for semantic boundaries"""
        chunks = []
        
        try:
            parser = self.parsers[language]
            tree = parser.parse(bytes(code, 'utf-8'))
            
            # Extract functions, classes, and other top-level constructs
            root_node = tree.root_node
            
            for child in root_node.children:
                if child.type in ['function_definition', 'class_definition', 'function_declarator']:
                    start_byte = child.start_byte
                    end_byte = child.end_byte
                    
                    chunk_content = code[start_byte:end_byte]
                    
                    chunks.append({
                        'content': chunk_content,
                        'content_type': 'code',
                        'chunk_index': len(chunks),
                        'metadata': {
                            'language': language,
                            'node_type': child.type,
                            'start_line': child.start_point[0],
                            'end_line': child.end_point[0]
                        }
                    })
            
            # If no semantic chunks found, fall back to simple chunking
            if not chunks:
                code_chunks = self._chunk_text(code, 'code')
                for i, chunk_text in enumerate(code_chunks):
                    chunks.append({
                        'content': chunk_text,
                        'content_type': 'code',
                        'chunk_index': i,
                        'metadata': {'language': language}
                    })
        
        except Exception as e:
            self.logger.warning(f"Error in semantic chunking: {e}")
            # Fall back to simple chunking
            code_chunks = self._chunk_text(code, 'code')
            for i, chunk_text in enumerate(code_chunks):
                chunks.append({
                    'content': chunk_text,
                    'content_type': 'code',
                    'chunk_index': i,
                    'metadata': {'language': language}
                })
        
        return chunks
    
    def _chunk_python_with_ast(self, code: str) -> List[Dict[str, Any]]:
        """Chunk Python code using AST"""
        chunks = []
        
        try:
            tree = ast.parse(code)
            
            for node in ast.walk(tree):
                if isinstance(node, (ast.FunctionDef, ast.ClassDef, ast.AsyncFunctionDef)):
                    # Get the source code for this node
                    start_line = node.lineno - 1
                    end_line = node.end_lineno if hasattr(node, 'end_lineno') else start_line + 10
                    
                    lines = code.split('\n')
                    chunk_content = '\n'.join(lines[start_line:end_line])
                    
                    chunks.append({
                        'content': chunk_content,
                        'content_type': 'code',
                        'chunk_index': len(chunks),
                        'metadata': {
                            'language': 'python',
                            'node_type': type(node).__name__,
                            'name': node.name,
                            'start_line': start_line,
                            'end_line': end_line
                        }
                    })
        
        except SyntaxError as e:
            self.logger.warning(f"Python syntax error, falling back to simple chunking: {e}")
            # Fall back to simple chunking
            code_chunks = self._chunk_text(code, 'code')
            for i, chunk_text in enumerate(code_chunks):
                chunks.append({
                    'content': chunk_text,
                    'content_type': 'code',
                    'chunk_index': i,
                    'metadata': {'language': 'python'}
                })
        
        return chunks
    
    def _chunk_matlab_functions(self, code: str) -> List[Tuple[str, str]]:
        """Chunk MATLAB code by functions"""
        chunks = []
        
        # Simple pattern matching for MATLAB functions
        function_pattern = r'^function\s+.*?^end\s*
        matches = re.finditer(function_pattern, code, re.MULTILINE | re.DOTALL)
        
        for match in matches:
            func_code = match.group(0)
            # Extract function name
            func_line = func_code.split('\n')[0]
            func_name_match = re.search(r'function.*?(\w+)\s*\(', func_line)
            func_name = func_name_match.group(1) if func_name_match else 'unknown'
            
            chunks.append((func_code, func_name))
        
        # If no functions found, treat as script
        if not chunks:
            chunks.append((code, 'script'))
        
        return chunks
    
    def _split_sql_statements(self, sql_content: str) -> List[str]:
        """Split SQL content into individual statements"""
        # Simple splitting by semicolon (could be improved)
        statements = []
        current_statement = ""
        
        for line in sql_content.split('\n'):
            line = line.strip()
            if line and not line.startswith('--'):  # Skip comments
                current_statement += line + '\n'
                
                if line.endswith(';'):
                    statements.append(current_statement.strip())
                    current_statement = ""
        
        # Add any remaining statement
        if current_statement.strip():
            statements.append(current_statement.strip())
        
        return statements
    
    def _identify_sql_statement_type(self, statement: str) -> str:
        """Identify the type of SQL statement"""
        statement_upper = statement.upper().strip()
        
        if statement_upper.startswith('SELECT'):
            return 'SELECT'
        elif statement_upper.startswith('INSERT'):
            return 'INSERT'
        elif statement_upper.startswith('UPDATE'):
            return 'UPDATE'
        elif statement_upper.startswith('DELETE'):
            return 'DELETE'
        elif statement_upper.startswith('CREATE'):
            return 'CREATE'
        elif statement_upper.startswith('ALTER'):
            return 'ALTER'
        elif statement_upper.startswith('DROP'):
            return 'DROP'
        else:
            return 'OTHER'
    
    # Utility methods
    
    def find_files(self, 
                   directory: str, 
                   recursive: bool = True,
                   file_patterns: List[str] = None) -> List[str]:
        """Find files to process in a directory"""
        directory = Path(directory)
        
        if not directory.exists():
            self.logger.error(f"Directory not found: {directory}")
            return []
        
        files = []
        supported_extensions = set(self.processors.keys())
        
        if recursive:
            pattern = "**/*"
        else:
            pattern = "*"
        
        for file_path in directory.glob(pattern):
            if file_path.is_file():
                # Check if file extension is supported
                if file_path.suffix.lower() in supported_extensions:
                    # Apply file patterns if provided
                    if file_patterns:
                        if any(pattern in str(file_path) for pattern in file_patterns):
                            files.append(str(file_path))
                    else:
                        files.append(str(file_path))
        
        self.logger.info(f"Found {len(files)} supported files in {directory}")
        return files
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions"""
        return list(self.processors.keys())
    
    def get_file_stats(self, file_path: str) -> Dict[str, Any]:
        """Get basic file statistics"""
        try:
            path = Path(file_path)
            stats = path.stat()
            
            return {
                'size': stats.st_size,
                'modified': stats.st_mtime,
                'created': stats.st_ctime,
                'extension': path.suffix.lower(),
                'name': path.name,
                'directory': str(path.parent)
            }
        except Exception as e:
            self.logger.error(f"Error getting file stats for {file_path}: {e}")
            return {}
    
    async def validate_file(self, file_path: str) -> Dict[str, Any]:
        """Validate if a file can be processed"""
        validation = {
            'valid': False,
            'file_exists': False,
            'supported_format': False,
            'readable': False,
            'size_ok': False,
            'issues': []
        }
        
        try:
            path = Path(file_path)
            
            # Check if file exists
            if not path.exists():
                validation['issues'].append('File does not exist')
                return validation
            validation['file_exists'] = True
            
            # Check if format is supported
            if path.suffix.lower() not in self.processors:
                validation['issues'].append(f'Unsupported file format: {path.suffix}')
                return validation
            validation['supported_format'] = True
            
            # Check if file is readable
            try:
                with open(path, 'rb') as f:
                    f.read(1024)  # Try to read first 1KB
                validation['readable'] = True
            except Exception as e:
                validation['issues'].append(f'File not readable: {e}')
                return validation
            
            # Check file size (max 100MB)
            file_size = path.stat().st_size
            if file_size > 100 * 1024 * 1024:
                validation['issues'].append(f'File too large: {file_size / (1024*1024):.1f}MB')
                return validation
            validation['size_ok'] = True
            
            validation['valid'] = True
            
        except Exception as e:
            validation['issues'].append(f'Validation error: {e}')
        
        return validation